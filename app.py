import os
import re
import difflib
import json
from flask import Flask, request, render_template, jsonify, send_file, url_for
from dotenv import load_dotenv
from zipfile import ZipFile
#import openai
import requests
import time
import datetime

from docx import Document
from docx.shared import Pt

# Import vlastných modulov (Telnet komunikácia, konfigurácia, report)
from src.communication.Communication import TelnetCommunication
from src.core.config import TypeOfConnection
#from config_diff_evaluator import clear_docx_content_in_place

# Definícia priečinka na nahrávanie
UPLOAD_FOLDER = "USER_REPORTS"



load_dotenv()

OPENWEBUI_API_URL = "http://pes.kis.fri.uniza.sk/api/chat/completions"
API_KEY = os.getenv("OPENAI_API_KEY")
AI_MODEL = "deepseek-r1:7b"

DEVICE_MANUFACTURERS = {
    "cisco": {
        "disable_paging": "terminal length 0",
        "show_config": "show running-config",
        "prompt_end": "#"
    },
    "mikrotik": {
        "disable_paging": None,
        "show_config": "/export",
        "prompt_end": "]"
    }
}


app = Flask(__name__)

UPLOAD_FILES_FOLDER = "USER_UPLOAD_FILES"
os.makedirs(UPLOAD_FILES_FOLDER, exist_ok=True)
app.config['USER_UPLOAD_FILES_FOLDER'] = UPLOAD_FILES_FOLDER
# STATIC_CONFIG_DIR = os.path.join(app.static_folder, "CONFIGS")
# os.makedirs(STATIC_CONFIG_DIR, exist_ok=True
STATIC_CONFIG_DIR = "USER_CONFIGS"
os.makedirs(STATIC_CONFIG_DIR, exist_ok=True)
app.config['USER_CONFIGS'] = STATIC_CONFIG_DIR
##############################################################################################################
# Funkcia na kontrolu povolených súborov
# Vráti True, ak je súbor povolený
# (má príponu .txt, .cfg, .conf)
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'txt', 'cfg', 'conf'}
##############################################################################################################
# Funkcia na odstránenie ANSI kódov z textu
def strip_ansi_codes(text: str) -> str:
    ansi_escape = re.compile(r'\x1B\[[0-?]*[ -/]*[@-~]')
    return ansi_escape.sub('', text)
##############################################################################################################
# Funkcia na odstránenie prázdnych riadkov a medzier
# z riadkov, ktoré sú v zozname
# a vrátenie ich ako zoznam
def flatten_diffs(entries) -> list:
    flat = []
    for entry in entries:
        for line in entry.splitlines():
            line = line.strip()
            if line:
                flat.append(line)
    print(f"flatten_diffs({entries}) -> {flat}")
    return flat
##############################################################################################################
# Funkcia na odstránenie všetkých riadkov pred /export
# a ponechanie len tých, ktoré sú za ním
# (pre Mikrotik export)
def trim_to_export(config: str) -> str:
    clean = strip_ansi_codes(config)
    lines = clean.splitlines()

    # nájdeme index posledného riadku, ktorý kdekoľvek obsahuje "/export"
    export_idxs = [i for i, line in enumerate(lines) if "/export" in line]
    if not export_idxs:
        # žiadny /export v texte => vrátime celý "clean"
        return clean

    start = export_idxs[-1] + 1  # +1 aby sme vylúčili aj ten riadok s príkazom
    tail = lines[start:]

    # prípadne ukončiť pred promptom (napr. "[admin@R1] >")
    end = len(tail)
    for i, line in enumerate(tail):
        text = line.strip()
        if text.startswith("[") and text.endswith(">"):
            end = i
            break

    # vyberieme pásmo a odfiltrujeme prázdne riadky
    extracted = tail[:end]
    filtered = [l for l in extracted if l.strip()]

    return "\n".join(filtered)

##############################################################################################################
# Funkcia na pripojenie k zariadeniu cez Telnet
# a získanie konfigurácie
# Vstup: IP adresa, port, užívateľské meno, heslo
# Výstup: konfigurácia ako reťazec
# Ak sa pripojenie nepodarí, vráti chybovú správu
# Ak sa pripojenie podarí, vráti konfiguráciu ako reťazec
def get_configuration(host: str, port: int, user: str, password: str):
    try:
        print(f"Pripájam sa k zariadeniu {host}:{port}...")
        telnet_comm = TelnetCommunication()
        telnet_comm.initConnection(host, user, password, port, TypeOfConnection.TELNET)
        telnet_comm.getRoot(user, password)

        manufacturer = getattr(telnet_comm, "manufacturer", "cisco")  # predvolený cisco
        print(f"Detekovaný výrobca: {manufacturer}")

        manufacturer_settings = DEVICE_MANUFACTURERS.get(manufacturer, DEVICE_MANUFACTURERS["cisco"])

        paging_command = manufacturer_settings["disable_paging"]
        show_command = manufacturer_settings["show_config"]
        prompt_end = manufacturer_settings["prompt_end"]

        if manufacturer == "mikrotik":
            telnet_comm.tConnection.write(show_command.encode('ascii') + b"\r\n")
            time.sleep(3)  # niekedy treba 2–3 sekundy
            raw = telnet_comm.tConnection.read_very_eager().decode('ascii')
            output = trim_to_export(raw)
        else:
            if paging_command:
                telnet_comm.tConnection.write(paging_command.encode('ascii') + b"\r\n")
                time.sleep(1)
                telnet_comm.tConnection.read_very_eager()  # vyčisti buffer
            telnet_comm.tConnection.write(show_command.encode('ascii') + b"\r\n")
            output = telnet_comm.tConnection.read_until(prompt_end.encode('ascii'), timeout=60).decode('ascii')

        
        telnet_comm.closeConnection()
        print(f"[DEBUG] Výstup z MikroTik:\n{output}")
        return output
    except Exception as e:
        return f"Error: {str(e)}"

##############################################################################################################
# Funkcia na extrakciu názvu zariadenia z konfigurácie
# Vstup: konfigurácia ako reťazec
# Výstup: názov zariadenia ako reťazec
# Ak sa názov zariadenia nedá získať, vráti "Neznáme_zariadenie"
def extract_hostname(config: str) -> str:
    # 1) odstrániť prípadné ANSI kódy, ak ich máte v texte
    ansi_escape = re.compile(r'\x1B\[[0-?]*[ -/]*[@-~]')
    clean = ansi_escape.sub('', config)

    # 2) skúsiť Cisco-style "hostname X"
    for line in clean.splitlines():
        m = re.match(r'^\s*hostname\s+(\S+)', line, re.IGNORECASE)
        if m:
            return m.group(1)

    # 3) skúsiť Mikrotik-style "user@HOST"
    #    zachytí admin@R1>, admin@R1:~#, …
    m = re.search(r'\b[\w-]+@([\w-]+)', clean)
    if m:
        return m.group(1)

    # 4) fallback: "set name=XYZ" v Mikrotik exporte
    m = re.search(r'\bname\s*=\s*(\S+)', clean, re.IGNORECASE)
    if m:
        return m.group(1)

    return "Neznáme_zariadenie"
##############################################################################################################
# Funkcia na výpočet podobnosti medzi dvoma konfiguráciami
# Vstup: dve konfigurácie ako reťazce
# Výstup: podobnosť ako percento (float)
# Ak sú obidve konfigurácie prázdne, vráti 100%
def calculate_similarity(config1: str, config2: str) -> float:
    # zbavíme sa prázdnych a čisto medzerových riadkov, orežeme odsadenia
    s1 = {line.strip() for line in config1.splitlines() if line.strip()}
    s2 = {line.strip() for line in config2.splitlines() if line.strip()}
    if not s1 and not s2:
        return 100.0
    matched = s1.intersection(s2)
    unioned = s1.union(s2)
    return round(len(matched) / len(unioned) * 100, 2)

##############################################################################################################
# Funkcia na zvýraznenie rozdielov medzi dvoma konfiguráciami
# Vstup: konfigurácia ako reťazec, zoznam chýbajúcich riadkov, zoznam pridaných riadkov
# Výstup: zvýraznená konfigurácia ako reťazec
# Ak je riadok v "order", zvýrazníme ho oranžovo (má prioritu)
# Inak pre referenčnú konfiguráciu: ak je v missing, zvýrazníme modrou
# Ale pre porovnávaciu konfiguráciu: ak je v extra, zvýrazníme červenou
def highlight_differences_multiple(config: str, missing: list, extra: list, order: list,
                                   missing_color: str = "blue", extra_color: str = "red", order_color: str = "orange"):
    highlighted_config = []
    # Normalizácia rozdielov pre každú kategóriu
    missing_norm = [re.sub(r'\s+', ' ', diff.strip()) for diff in missing]
    extra_norm   = [re.sub(r'\s+', ' ', diff.strip()) for diff in extra]
    order_norm   = [re.sub(r'\s+', ' ', diff.strip()) for diff in order]
    
    for index, line in enumerate(config.splitlines(), start=1):
        normalized_line = re.sub(r'\s+', ' ', line.strip())
        color = None
        # Ak je riadok v "order", zvýrazníme ho oranžovo (má prioritu)
        if normalized_line in order_norm:
            color = order_color
        # Inak pre referenčnú konfiguráciu: ak je v missing, zvýrazníme modrou
        elif normalized_line in missing_norm:
            color = missing_color
        # Ale pre porovnávaciu konfiguráciu: ak je v extra, zvýrazníme červenou
        elif normalized_line in extra_norm:
            color = extra_color

        if color:
            highlighted_line = f'<span style="color: {color}; font-weight: bold;">{line}</span>'
        else:
            highlighted_line = line
        line_number = f'<span style="display: inline-block; width: 30px; text-align: right; color: #888;">{index}</span>'
        highlighted_config.append(f'{line_number} {highlighted_line}')
    
    return "<br>".join(highlighted_config)
##############################################################################################################
# Funkcia na porovnanie dvoch konfigurácií
# Vstup: dve konfigurácie ako reťazce
# Výstup: slovník s kľúčmi "missing", "added", "order"
# a hodnotami ako zoznamy riadkov
# Ak sa nájde rozdiel, vráti ho ako slovník
# Ak sa nájde chyba, vráti ju ako slovník s kľúčom "error"
def     analyze_config_differences_cross_platform(ref_config: str, comp_config: str):
    """
    Cisco + Mikrotik porovnanie:
    - missing / added riadky
    - poradie príkazov (order-sensitive)
    - zámenné IP adresy medzi rozhraniami sú spracované ako chýbajúce/pridané
    """

    def normalize_line(line: str) -> str:
        line = re.sub(r"!.*", "", line)
        line = re.sub(r"#.*", "", line)
        return re.sub(r"\s+", " ", line.strip().lower())

    def is_order_sensitive(line: str) -> bool:
        keywords = ["access-list", "route-map", "policy-map", "class-map", "/ip firewall", "/queue", "/interface bridge"]
        return any(line.lower().startswith(k) for k in keywords)

    def extract_interface_ips(lines: list[str]) -> dict:
        """
        Extrahuje mapovanie rozhranie → IP adresa (podpora pre Cisco aj Mikrotik)
        """
        interfaces = {}
        current_iface = None

        for line in lines:
            # Cisco
            if line.lower().startswith("interface"):
                current_iface = line.split()[1]
            elif current_iface and "ip address" in line.lower():
                match = re.search(r"ip address (\S+ \S+)", line)
                if match:
                    interfaces[current_iface] = match.group(1)

            # MikroTik
            elif "/ip address add" in line:
                iface_match = re.search(r"interface=(\\S+)", line)
                ip_match = re.search(r"address=(\\S+)", line)
                if iface_match and ip_match:
                    interfaces[iface_match.group(1)] = ip_match.group(1)

        return interfaces

    ref_lines = [line for line in ref_config.splitlines() if line.strip()]
    comp_lines = [line for line in comp_config.splitlines() if line.strip()]
    ref_norm = [normalize_line(line) for line in ref_lines]
    comp_norm = [normalize_line(line) for line in comp_lines]

    matcher = difflib.SequenceMatcher(None, ref_norm, comp_norm)
    missing, added, order = [], [], []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        elif tag == "replace":
            for line in ref_lines[i1:i2]:
                if is_order_sensitive(line):
                    order.append(line)
                else:
                    missing.append(line)
            for line in comp_lines[j1:j2]:
                if is_order_sensitive(line):
                    order.append(line)
                else:
                    added.append(line)
        elif tag == "delete":
            missing.extend(ref_lines[i1:i2])
        elif tag == "insert":
            added.extend(comp_lines[j1:j2])

    # Zámenné IP adresy: ak je IP adresa na inom rozhraní ako má byť, považuj to za chýbajúcu + pridanú
    ref_ips = extract_interface_ips(ref_lines)
    comp_ips = extract_interface_ips(comp_lines)

    for iface, ip in ref_ips.items():
        found = any(c_ip == ip for c_ip in comp_ips.values())
        if iface not in comp_ips or (iface in comp_ips and comp_ips[iface] != ip and found):
            # IP adresa je rovnaká, ale priradená nesprávne → chyba poradia
            missing.append(f"{iface} missing IP {ip}")
            # nájdi, kde sa nachádza a zaraď ako extra
            for c_iface, c_ip in comp_ips.items():
                if c_ip == ip and c_iface != iface:
                    added.append(f"{c_iface} has unexpected IP {c_ip}")

    return {
        "missing": missing,
        "added": added,
        "order": order
    }

##############################################################################################################
# Funkcia na kontrolu poradia príkazov pomocou AI
# Vstup: názov zariadenia, porovnávaná konfigurácia, referenčná konfigurácia
# model AI, API kľúč (ak je potrebný), zdroj API (školský alebo OpenAI)
# Výstup: zoznam príkazov, ktoré sú v nesprávnom poradí
# Ak sa nájde chyba, vráti ju ako prázdny zoznam
def ai_check_command_order(device_name: str, compare_config: str, correct_config: str, model, api_key=None, api_source="school"):
    """
    Skontroluje cez AI, či je poradie príkazov správne.
    """

    prompt = (
        f"For the device {device_name}, compare the two network configurations listed below:\n\n"
        "---------------------------------------------------------------\n"
        f"Configuration 1 (reference):\n{correct_config}\n\n"
        "---------------------------------------------------------------\n"
        f"Configuration 2 (comparison):\n{compare_config}\n\n"
        "---------------------------------------------------------------\n"
        "Instructions:\n"
        "1. Focus only on sections where the order of commands matters (such as access-lists, route-maps, policy-maps, class-maps, etc.).\n"
        "2. Identify if the commands in these sections appear in a different order between Configuration 1 and Configuration 2.\n"
        "3. If any commands are found to be in the wrong order compared to the reference configuration, list them under the key \"order\".\n"
        "4. If all order-sensitive sections are identical in both content and order, return an empty list for \"order\".\n\n"
        "Return the result strictly in valid JSON format, like this:\n"
        "{\n"
        "  \"order\": [\"command1\", \"command2\", \"command3\"]\n"
        "}\n\n"
        "DO NOT add any additional explanation, text, or formatting outside of the JSON object.\n"
        "Return only the JSON object."
    )

    headers = { "Content-Type": "application/json" }

    if api_source == "openai":
        url = "https://api.openai.com/v1/chat/completions"
        headers["Authorization"] = f"Bearer {api_key}"
    else:
        url = OPENWEBUI_API_URL
        headers["Authorization"] = f"Bearer {API_KEY}"

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a network configuration expert. Respond only with valid JSON."},
            {"role": "user", "content": prompt}
        ]
    }

    response = requests.post(OPENWEBUI_API_URL, headers=headers, data=json.dumps(payload))

    if response.status_code == 200:
        try:
            result_text = response.json()["choices"][0]["message"]["content"]
            result_text_clean = result_text.strip()
            json_match = re.search(r"\{.*?\}", result_text_clean, re.DOTALL)
            if json_match:
                json_text = json_match.group(0)
                result_json = json.loads(json_text)
                return result_json.get("order", [])
            else:
                return []
        except (json.JSONDecodeError, KeyError):
            return []
    else:
        return []
##############################################################################################################
# Funkcia na porovnanie konfigurácií pomocou AI
# Vstup: názov zariadenia, porovnávaná konfigurácia, referenčná konfigurácia
# model AI, API kľúč (ak je potrebný), zdroj API (školský alebo OpenAI)
# Výstup: slovník s kľúčmi "missing", "added", "order" a "report"
# a hodnotami ako zoznamy riadkov
# Ak sa nájde chyba, vráti ju ako slovník s kľúčom "error"
# a hodnotou chybovej správy
def ai_compare_configurations(device_name, compare_config, correct_config, model, api_key=None, api_source="school"):

    print("===[ AI porovnávanie konfigurácií - nacitanie promptu ]===")
    prompt = (
        f'You are a senior network engineer tasked with comparing two configurations for the device {device_name}. '
        'These configurations come from Cisco or MikroTik devices.\n'
        'Your objective is to analyze them and return precise functional differences. Use the following structure and rules:\n\n'
        '========= Configuration 1 (Reference) =========\n'
        f'{correct_config}\n'
        '==============================================\n'
        '========= Configuration 2 (Compared) =========\n'
        f'{compare_config}\n'
        '==============================================\n\n'
        'Instructions:'
        '1. Identify lines that exist in the reference configuration but are completely missing in the compared configuration. Include them under `"missing"`.'
        '2. Identify lines that exist only in the compared configuration (not in the reference). Include them under `"added"`.'
        '3. If a command is the same but appears in a different place (especially inside order-sensitive blocks like access-lists, route-maps, policy-maps, MikroTik firewall rules), include it under `"order"`.'
        '4. If there is a functional difference between two lines (e.g. IP address, mask, the presence of `no`, wrong interface, etc.), include the reference line in `"missing"` and the incorrect line in `"added"`.'
        '5. Do NOT treat small formatting changes (e.g. indentation, spacing) as differences.'
        '6. If an IP address is assigned to the wrong interface (swapped), treat it as one missing and one added line.'
        '7. Generate a concise technical report in English under the key `"report"`. The report should:'
        '   - Describe the core configuration differences'
        '   - Mention any security or operational risks'
        '   - Suggest what should be corrected'
        'Return ONLY the following valid JSON object:'
        '```json'
        '{'
        '  "missing": ["..."],'
        '  "added": ["..."],'
        '  "order": ["..."],'
        '  "report": "..."'
        '}'
        'Return ONLY JSON !!!'

    )
    
    print("===[ AI porovnávanie konfigurácií - nacitanie headeru ]===")
    headers = { "Content-Type": "application/json" }

    if api_source == "openai":
        url = "https://api.openai.com/v1/chat/completions"
        headers["Authorization"] = f"Bearer {api_key}"
    else:
        url = OPENWEBUI_API_URL
        headers["Authorization"] = f"Bearer {API_KEY}"
    print("===[ AI porovnávanie konfigurácií - nacitanie payloadu ]===")
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "You are a network configuration expert. Respond only with valid JSON."},
            {"role": "user", "content": prompt}
        ]
    }
    print("===[ AI porovnávanie konfigurácií - nacitanie requestu ]===")
    print(f"Payload: {json.dumps(payload, indent=2)}")
    print(f"URL: {url}")
    print(f"Headers: {headers}")
    response = requests.post(url, headers=headers, data=json.dumps(payload))
    print("===[ AI porovnávanie konfigurácií - nacitanie response ]===")
    if response.status_code == 200:
        try:
            result_text = response.json()["choices"][0]["message"]["content"]
            print("=== AI RAW RESPONSE ===")
            print(result_text)

            # Pokus č. 1 – načítanie ako čistý JSON
            try:
                result_json = json.loads(result_text)
            except json.JSONDecodeError:
                # Pokus č. 2 – extrakcia JSON objektu zo surového textu
                json_match = re.search(r"\{.*\}", result_text, re.DOTALL)
                if json_match:
                    try:
                        result_json = json.loads(json_match.group(0))
                    except Exception as e:
                        print(f"Second-level JSON parse failed: {str(e)}")
                        return {
                            "error": f"Second-level JSON parse failed: {str(e)}",
                            "raw_output": result_text
                        }
                else:
                    print("No JSON object found in AI response.")
                    return {
                        "error": "Could not extract JSON object from AI response.",
                        "raw_output": result_text
                    }

            # Úspešný parse
            return {
                "missing": result_json.get("missing", []),
                "added": result_json.get("added", []),
                "order": result_json.get("order", []),
                "report": result_json.get("report", "")
            }

        except Exception as e:
            print(f"Error parsing AI response: {str(e)}")
            return {"error": f"Unexpected error while parsing AI response: {str(e)}"}

    else:
        print(f"API error: {response.status_code} {response.text}")
        return {"error": f"API error: {response.status_code} {response.text}"}

##############################################################################################################
# Funkcia na vygenerovanie unikátneho názvu súboru pre report
# Vstup: žiadny
# Výstup: názov súboru ako reťazec
# Vytvorí priečinok, ak neexistuje
# a vygeneruje názov súboru vo formáte "Report_YYYYMMDD_X.docx"
def generate_report_filename():
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    os.makedirs(f"{UPLOAD_FOLDER}/{today}", exist_ok=True)

    # Find the next available index
    index = 1
    while os.path.exists(f"{UPLOAD_FOLDER}/{today}/Report_{today}_{index}.docx"):
        index += 1

    return f"{UPLOAD_FOLDER}/{today}/Report_{today}_{index}.docx"
##############################################################################################################
# Funkcia na vygenerovanie Word reportu
# Vstup: správy ako slovník, cesta k súboru ako reťazec
# Výstup: žiadny
# Vytvorí Word dokument a uloží ho na zadanú cestu
# Pre každé zariadenie pridá nadpis, report, pridané príkazy, chýbajúce príkazy a zlé poradie
# Ak sa správa nenájde, pridá prázdny riadok
def generate_word_report(reports: dict, file_path: str):
    doc = Document()
    print("Generujem Word report pre všetky zariadenia...")
    for device, diff in reports.items():
        print(f"Spracovávam report pre zariadenie: {device}")
        doc.add_heading(device, level=1)
        
        doc.add_heading("Report:", level=2)
        report_text = diff.get("report", "")
        print(f"Report pre {device}: {report_text}")
        doc.add_paragraph(report_text if report_text else "Žiadny report.")
        
        doc.add_heading("Pridané príkazy:", level=2)
        added = diff.get("added", [])
        print(f"Pridané príkazy pre {device}: {added}")
        if added:
            for cmd in added:
                doc.add_paragraph(cmd, style='List Bullet')
        else:
            doc.add_paragraph("Žiadne pridané príkazy.")
        
        doc.add_heading("Chýbajúce príkazy:", level=2)
        missing = diff.get("missing", [])
        print(f"Chýbajúce príkazy pre {device}: {missing}")
        if missing:
            for cmd in missing:
                doc.add_paragraph(cmd, style='List Bullet')
        else:
            doc.add_paragraph("Žiadne chýbajúce príkazy.")
        
        doc.add_heading("Zlé poradie:", level=2)
        order = diff.get("order", [])
        print(f"Zlé poradie pre {device}: {order}")
        if order:
            for cmd in order:
                doc.add_paragraph(cmd, style='List Bullet')
        else:
            doc.add_paragraph("Poradie je v poriadku.")
        
        doc.add_paragraph("")  # medzera medzi zariadeniami
    doc.save(file_path)
    print(f"Word report bol uložený na: {file_path}")
##############################################################################################################
# Funkcia na porovnanie konfigurácií zo súborov
# Vstup: referenčné a porovnávané súbory ako zoznamy, režim porovnania
# model AI, API kľúč (ak je potrebný)
# Výstup: správa ako slovník
@app.route('/compare_files', methods=['POST'])
def compare_files():
    try:
        print
        # Získaj súbory
        ref_files = request.files.getlist('ref_files')
        comp_files = request.files.getlist('comp_files')
        compare_mode = request.form.get("compare_mode", "school_ai_line_by_line")
        selected_model = request.form.get("model", AI_MODEL)
        api_key_override = request.form.get("api_key", None)

        api_source = "openai" if compare_mode == "chatgpt_ai" else "school"

        print("===[ FRONTEND INPUT ]===")
        print(f"Compare mode: {compare_mode}")
        print(f"Selected model: {selected_model}")
        print(f"API source: {api_source}")
        print(f"API key override provided: {'YES' if api_key_override else 'NO'}")


        if len(ref_files) != len(comp_files):
            return jsonify({"error": "Počet referenčných a porovnávaných súborov sa musí zhodovať."}), 400

        reports = {}
        correctoutput = ""
        compareoutput = ""
        similarity_score = 0

        report_path = generate_report_filename()

        for i, (ref_file, comp_file) in enumerate(zip(ref_files, comp_files)):
            if not (allowed_file(ref_file.filename) and allowed_file(comp_file.filename)):
                continue

            ref_path = os.path.join(UPLOAD_FILES_FOLDER, f"ref_{i}_{ref_file.filename}")
            comp_path = os.path.join(UPLOAD_FILES_FOLDER, f"comp_{i}_{comp_file.filename}")
            ref_file.save(ref_path)
            comp_file.save(comp_path)

            with open(ref_path, "r", encoding="utf-8", errors="ignore") as f:
                correct_config = f.read()
            with open(comp_path, "r", encoding="utf-8", errors="ignore") as f:
                compare_config = f.read()

            similarity_score += calculate_similarity(correct_config, compare_config)
            name = extract_hostname(correct_config)

            
            if compare_mode == "line_by_line":
                differences = analyze_config_differences_cross_platform(correct_config, compare_config)
                print("===[ konfigurácie porovnané - riadkovo ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"Comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = "Line-by-line comparison completed without AI assistance."
            elif compare_mode == "school_ai_line_by_line":
                differences = analyze_config_differences_cross_platform(correct_config, compare_config)
                print("===[ konfigurácie porovnané - riadkovo + školská AI ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"AI comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = differences.get("report", "")
                order_lines += ai_check_command_order(name, correct_config, compare_config, selected_model, api_key_override)
                report_text = f"Line-by-line comparison with AI for order analysis. {len(order_lines)} order-related issues detected."
            elif compare_mode in ("school_ai", "chatgpt_ai"):
                differences = ai_compare_configurations(name, compare_config, correct_config, selected_model, api_key_override, api_source)
                print("===[ konfigurácie porovnané - AI ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"AI comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = differences.get("report", "")
            else:
                missing_lines, added_lines, order_lines = [], [], []
                report_text = "Unsupported comparison mode."

                
            differences = {
                "missing": missing_lines,
                "added": added_lines,
                "order": order_lines,
                "report": report_text
            }
            
            print("Rozdiely:", differences)
            
            # HTML zvýraznenie pre web
            print("HTML zvýraznenie pre web")

            # … v compare_files(), tesne pred HTML zvýraznením:
            # pôvodne si mal:
            # highlighted_correct = highlight_differences_multiple(correct_config, missing_lines, [], order_lines)
            # highlighted_compare = highlight_differences_multiple(compare_config, [], added_lines, order_lines)

            # namiesto toho vlož:
            missing_flat = flatten_diffs(missing_lines)
            added_flat   = flatten_diffs(added_lines)
            order_flat   = flatten_diffs(order_lines)

            highlighted_correct = highlight_differences_multiple(
                correct_config,
                missing_flat,
                [],           # extra = []
                order_flat
            )
            highlighted_compare = highlight_differences_multiple(
                compare_config,
                [],           # missing = []
                added_flat,
                order_flat
            )

            print("pridanie hlaicky")
            correctoutput += f"<div><h4>{name} - referenčná konfigurácia</h4>{highlighted_correct}</div>"
            compareoutput += f"<div><h4>{name} - porovnávacia konfigurácia</h4>{highlighted_compare}</div>"

            # Uloženie reportu pre Word dokument
            print("Uloženie reportu pre Word dokument")
            reports[name] = {
                "missing": missing_lines,
                "added": added_lines,
                "order": order_lines,
                "report": report_text
            }
            print(f"Report pre {name}: {reports[name]}")

        print("reports:", reports)
        print("reports.keys():", reports.keys())
        print("report path:", report_path)
        # Generovanie Word reportu so štruktúrou: Názov zariadenia, Report, Pridané príkazy, Chýbajúce príkazy, Zlé poradie
        print("Generovanie Word reportu")
        generate_word_report(reports, report_path)
        
        report_filename_only = os.path.basename(report_path)
        similarity_score = round(similarity_score / len(ref_files), 2) if ref_files else 0
        print("Podobnosť konfigurácií:", similarity_score)
        print("Koniec porovnávania")
        result = {
            "correctoutput": correctoutput,
            "compareoutput": compareoutput,
            "similarity": similarity_score,
            "report": report_filename_only
            }
        return jsonify(result)
    except Exception as e:
        print(f"❌ Chyba v /compare_files: {str(e)}")
        return jsonify({"error": f"Chyba: {str(e)}"}), 500

##############################################################################################################
# Funkcia na porovnanie konfigurácií
# Vstup: referenčná konfigurácia, porovnávané konfigurácie
# Výstup: slovník s kľúčmi "missing", "added", "order" a "report"
# a hodnotami ako zoznamy riadkov
# Ak sa nájde chyba, vráti ju ako slovník s kľúčom "error"
# a hodnotou chybovej správy
@app.route('/compare', methods=['POST'])
def compare():
    try:
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        data = request.json
        if not data:
            return jsonify({"error": "Chýbajú vstupné údaje"}), 400

        correctoutput = ""
        compareoutput = ""
        reports = {}  # report pre Word dokument, kľúč = názov zariadenia
        similarity_score = 0
        devices = 0

        correct_router = data.get("correct_router")
        compare_routers = data.get("routers_with_errors", [])
        compare_mode = data.get("compare_mode", "school_ai_line_by_line")
        selected_model = data.get("model", AI_MODEL)
        api_key_override = data.get("api_key", None)
        api_source = "openai" if compare_mode == "chatgpt_ai" else "school"

        print("===[ FRONTEND INPUT ]===")
        print(f"Compare mode: {compare_mode}")
        print(f"Selected model: {selected_model}")
        print(f"API source: {api_source}")
        print(f"API key override provided: {'YES' if api_key_override else 'NO'}")


        if not correct_router or not correct_router.get("addresses"):
            return jsonify({"error": "Chýbajú údaje pre referenčnú konfiguráciu"}), 400
            
        report_path = generate_report_filename()
        #clear_docx_content_in_place(report_path)
        
        for correct_address, compare_router in zip(correct_router["addresses"], compare_routers):
            devices += 1
            correct_ip, correct_port = correct_address.split(":")
            compare_ip, compare_port = compare_router["address"].split(":")
            correct_config = get_configuration(correct_ip, int(correct_port),
                                                      correct_router["username"], correct_router["password"])
            compare_config = get_configuration(compare_ip, int(compare_port),
                                                      compare_router["username"], compare_router["password"])

            if not correct_config or not compare_config:
                correctoutput += f"<div style='color: red;'>Chyba pri získaní konfigurácie pre {correct_ip}</div>"
                continue
            
            print("===[ Porovnávam konfigurácie ]===")

            #print("correct_config:", correct_config)
            #print("compare_config:", compare_config)
            similarity_score += calculate_similarity(correct_config, compare_config)
            name = extract_hostname(correct_config)
            
            if compare_mode == "line_by_line":
                differences = analyze_config_differences_cross_platform(correct_config, compare_config)
                print("===[ konfigurácie porovnané - riadkovo ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"Comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = "Line-by-line comparison completed without AI assistance."
            elif compare_mode == "school_ai_line_by_line":
                differences = analyze_config_differences_cross_platform(correct_config, compare_config)
                print("===[ konfigurácie porovnané - riadkovo + školská AI ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"AI comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = differences.get("report", "")
                order_lines += ai_check_command_order(name, correct_config, compare_config, selected_model, api_key_override)
                report_text = f"Line-by-line comparison with AI for order analysis. {len(order_lines)} order-related issues detected."
            elif compare_mode in ("school_ai", "chatgpt_ai"):
                differences = ai_compare_configurations(name, compare_config, correct_config, selected_model, api_key_override, api_source)
                print("===[ konfigurácie porovnané - AI ]===")
                if "error" in differences:
                    print("⚠️ AI chyba:", differences["error"])
                    correctoutput += f"<div style='color: red;'>AI ERROR for {name}: {differences.get('error')}</div>"
                    correctoutput += f"<pre>{differences.get('raw_output', '')}</pre>"
                    missing_lines, added_lines, order_lines = [], [], []
                    report_text = f"AI comparison failed: {differences.get('error')}"
                else:
                    missing_lines = differences.get("missing", [])
                    added_lines = differences.get("added", [])
                    order_lines = differences.get("order", [])
                    report_text = differences.get("report", "")
            else:
                missing_lines, added_lines, order_lines = [], [], []
                report_text = "Unsupported comparison mode."

                
            differences = {
                "missing": missing_lines,
                "added": added_lines,
                "order": order_lines,
                "report": report_text
            }
            
            print("Rozdiely:", differences)
            
            # HTML zvýraznenie pre web
            print("HTML zvýraznenie pre web")

            # spláchnutie všetkých multi‐riadkových diff‐vstupov na jednotlivé riadky
            missing_flat = flatten_diffs(missing_lines)
            added_flat   = flatten_diffs(added_lines)
            order_flat   = flatten_diffs(order_lines)

            highlighted_correct = highlight_differences_multiple(
                correct_config,
                missing_flat,   # modré
                [],             # žiadne „extra“ tu
                order_flat      # oranžové
            )
            highlighted_compare = highlight_differences_multiple(
                compare_config,
                [],             # žiadne „missing“ tu
                added_flat,     # červené
                order_flat      # oranžové
            )

            print("pridanie hlaicky")
            correctoutput += f"<div><h4>{name} - referenčná konfigurácia</h4>{highlighted_correct}</div>"
            compareoutput += f"<div><h4>{name} - porovnávacia konfigurácia</h4>{highlighted_compare}</div>"

            # Uloženie reportu pre Word dokument
            print("Uloženie reportu pre Word dokument")
            reports[name] = {
                "missing": missing_lines,
                "added": added_lines,
                "order": order_lines,
                "report": report_text
            }
            print(f"Report pre {name}: {reports[name]}")

        print("reports:", reports)
        print("reports.keys():", reports.keys())
        print("report path:", report_path)
        # Generovanie Word reportu so štruktúrou: Názov zariadenia, Report, Pridané príkazy, Chýbajúce príkazy, Zlé poradie
        print("Generovanie Word reportu")
        generate_word_report(reports, report_path)
        
        report_filename_only = os.path.basename(report_path)
        similarity_score = round(similarity_score / devices, 2) if devices > 0 else 0
        print("Podobnosť konfigurácií:", similarity_score)
        print("Koniec porovnávania")
        result = {
            "correctoutput": correctoutput,
            "compareoutput": compareoutput,
            "similarity": similarity_score,
            "report": report_filename_only
            }
        return jsonify(result)

    except Exception as e:
        print(f"Chyba: {str(e)}")
        return jsonify({"error": f"Chyba: {str(e)}"}), 500
##############################################################################################################
@app.route('/compare_mixed', methods=['POST'])
def compare_mixed():
    try:
        # 1) Načítame uploadované súbory
        ref_files  = request.files.getlist("ref_files")
        comp_files = request.files.getlist("comp_files")

        # 2) Texty konfigurácií
        ref_configs, comp_configs = [], []

        # --- file vs. device ---
        if ref_files and not comp_files:
            # ref z nahraných súborov
            for f in ref_files:
                if not allowed_file(f.filename):
                    return jsonify(error=f"Nepodporovaný formát: {f.filename}"), 400
                ref_configs.append(f.read().decode("utf-8", errors="ignore"))
            # comp sťahované zo zariadení
            comp_mode = request.form.get("comp_mode", "download_device")
            if comp_mode == "download_device":
                user  = request.form.get("comp_user","").strip()
                pwd   = request.form.get("comp_pass","").strip()
                addrs = request.form.getlist("comp_addresses")
                if not user or not addrs:
                    return jsonify(error="Chýbajú údaje pre zariadenia."), 400
                for addr in addrs:
                    ip, port = addr.split(":")
                    comp_configs.append(get_configuration(ip, int(port), user, pwd))
            else:  # download_device_auth
                auths = request.form.getlist("comp_addresses_auth")
                for line in auths:
                    ipport, user, *rest = line.split(";")
                    pwd = rest[0] if rest else ""
                    ip, port = ipport.split(":")
                    comp_configs.append(get_configuration(ip, int(port), user, pwd))

        # --- device vs. file ---
        elif comp_files and not ref_files:
            # comp z nahraných súborov
            for f in comp_files:
                if not allowed_file(f.filename):
                    return jsonify(error=f"Nepodporovaný formát: {f.filename}"), 400
                comp_configs.append(f.read().decode("utf-8", errors="ignore"))
            # ref sťahované zo zariadení
            ref_mode = request.form.get("ref_mode", "download_device")
            if ref_mode == "download_device":
                user  = request.form.get("ref_user","").strip()
                pwd   = request.form.get("ref_pass","").strip()
                addrs = request.form.getlist("ref_addresses")
                if not user or not addrs:
                    return jsonify(error="Chýbajú údaje pre referenciu."), 400
                for addr in addrs:
                    ip, port = addr.split(":")
                    ref_configs.append(get_configuration(ip, int(port), user, pwd))
            else:  # download_device_auth
                auths = request.form.getlist("ref_addresses_auth")
                for line in auths:
                    ipport, user, *rest = line.split(";")
                    pwd = rest[0] if rest else ""
                    ip, port = ipport.split(":")
                    ref_configs.append(get_configuration(ip, int(port), user, pwd))

        else:
            return jsonify(error="Nepodporovaná kombinácia vstupov."), 400

        # 3) Porovnanie: obe polia už majú rovnakú dĺžku
        correctoutput = ""
        compareoutput = ""
        sim_score = 0.0
        reports = {}
        compare_mode    = request.form.get("compare_mode", "school_ai_line_by_line")
        selected_model  = request.form.get("model", AI_MODEL)
        api_key_override= request.form.get("api_key", None)
        api_source      = "openai" if compare_mode=="chatgpt_ai" else "school"

        for ref_cfg, comp_cfg in zip(ref_configs, comp_configs):
            sim_score += calculate_similarity(ref_cfg, comp_cfg)
            name = extract_hostname(ref_cfg)

            # vyhodnotenie podľa módu
            if compare_mode == "line_by_line":
                diffs = analyze_config_differences_cross_platform(ref_cfg, comp_cfg)
                report_text = "Line-by-line comparison"
            elif compare_mode == "school_ai_line_by_line":
                diffs = analyze_config_differences_cross_platform(ref_cfg, comp_cfg)
                diffs["order"] += ai_check_command_order(name, ref_cfg, comp_cfg,
                                                         selected_model, api_key_override)
                report_text = f"AI order analysis: {len(diffs['order'])} issues"
            else:  # school_ai alebo chatgpt_ai
                diffs = ai_compare_configurations(name, comp_cfg, ref_cfg,
                                                  selected_model, api_key_override,
                                                  api_source)
                report_text = diffs.get("report","")

            # zvýraznenie a HTML výstup
            missing = flatten_diffs(diffs.get("missing", []))
            added   = flatten_diffs(diffs.get("added",   []))
            order   = flatten_diffs(diffs.get("order",   []))

            highlighted_ref  = highlight_differences_multiple(ref_cfg,  missing, [],    order)
            highlighted_comp = highlight_differences_multiple(comp_cfg, [],      added, order)

            correctoutput += f"<div><h4>{name}</h4>{highlighted_ref}</div>"
            compareoutput += f"<div><h4>{name}</h4>{highlighted_comp}</div>"

            reports[name] = {
                "missing": diffs.get("missing", []),
                "added":   diffs.get("added",   []),
                "order":   diffs.get("order",   []),
                "report":  report_text
            }

        # 4) Word-report
        report_path = generate_report_filename()
        generate_word_report(reports, report_path)

        avg_sim = round(sim_score / len(ref_configs), 2) if ref_configs else 0.0
        return jsonify({
            "correctoutput": correctoutput,
            "compareoutput": compareoutput,
            "similarity": avg_sim,
            "report": os.path.basename(report_path)
        })

    except Exception as e:
        print("❌ Chyba v /compare_mixed:", e)
        return jsonify(error=str(e)), 500
##############################################################################################################
# Funkcia na stiahnutie reportu
# Vstup: názov súboru
# Výstup: súbor na stiahnutie
@app.route('/download_report', methods=['GET'])
def download_report():
    filename = request.args.get('file')
    file_path = os.path.abspath(os.path.join(UPLOAD_FOLDER, filename))  # Use absolute path

    print(f"Requested file: {filename}")
    print(f"Full absolute path: {file_path}")
    print(f"File exists: {os.path.exists(file_path)}")
    print(f"Is directory: {os.path.isdir(file_path)}")
    print(f"Is file: {os.path.isfile(file_path)}")

    if filename and os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)  # Use send_file instead of send_from_directory

    return "File not found", 404
##############################################################################################################
@app.route('/download-configs', methods=['POST'])
def download_configs():
    data      = request.json or {}
    addresses = data.get("addresses", [])
    username  = data.get("username", "")
    password  = data.get("password", "")

    if not addresses:
        return jsonify({"error": "Neboli zadané žiadne zariadenia."}), 400

    # Vymažeme staré súbory
    for f in os.listdir(STATIC_CONFIG_DIR):
        os.remove(os.path.join(STATIC_CONFIG_DIR, f))

    file_urls  = []
    full_text  = ""

    # Pre každý router vygenerujeme .txt
    for addr in addresses:
        ip, port = addr.split(":")
        cfg = get_configuration(ip, int(port), username, password)
        hostname = extract_hostname(cfg)
        fname    = f"{hostname}.txt"
        fpath    = os.path.join(STATIC_CONFIG_DIR, fname)

        with open(fpath, "w", encoding="utf-8") as f:
            f.write(cfg)

        # URL pre front-end (relatívna na /static/)
        file_urls.append(url_for('static', filename=f"configs/{fname}", _external=False))
        full_text += f"\n==== {hostname} ({ip}:{port}) ====\n{cfg}\n"

    bundle_url = None
    # Ak je viacero konfiguračných súborov, vytvor ZIP
    if len(file_urls) > 1:
        bundle_name = "configs_bundle.zip"
        bundle_path = os.path.join(STATIC_CONFIG_DIR, bundle_name)

        with ZipFile(bundle_path, 'w') as zf:
            for u in file_urls:
                local = os.path.join(STATIC_CONFIG_DIR, os.path.basename(u))
                zf.write(local, arcname=os.path.basename(local))

        bundle_url = url_for('static', filename=f"configs/{bundle_name}", _external=False)

    return jsonify({
        "configs":    full_text,
        "file_urls":  file_urls,
        "bundle_url": bundle_url
    })
##############################################################################################################
# Funkcia na stiahnutie ZIP súboru s konfiguráciami
# Vstup: žiadny
# Výstup: ZIP súbor na stiahnutie
@app.route('/download_bundle', methods=['GET'])
def download_bundle():
    bundle_name = 'configs_bundle.zip'
    # cesta ku ZIP-ku, ktorý si predtým vygeneroval/shutil.make_archive
    bundle_path = os.path.abspath(os.path.join('static', 'configs', bundle_name))

    if os.path.isfile(bundle_path):
        # send_file automaticky nastaví Content-Disposition: attachment
        return send_file(bundle_path, as_attachment=True)
    else:
        return "Bundle not found", 404
##############################################################################################################
# Funkcia na zobrazenie hlavnej stránky
# Vstup: žiadny
# Výstup: HTML stránka
@app.route('/')
def index():
    return render_template('index.html')
##############################################################################################################
# Funkcia na zobrazenie stránky s inštrukciami
# Vstup: žiadny
# Výstup: HTML stránka
@app.route('/download')
def download():
    return render_template('download_config.html')
##############################################################################################################
# Funkcia na zobrazenie stránky s inštrukciami pre porovnávanie
# Vstup: žiadny
# Výstup: HTML stránka
@app.route('/batch')
def batch():
    return render_template('batch.html')
##############################################################################################################
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5001))
    app.run(debug=True, host="0.0.0.0", port=port)
